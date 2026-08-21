"""
Full tex -> docx conversion (merge + preprocess + pandoc + typesetting polish)
Usage: python src/tex_to_docx.py
Output: paper_overleaf.docx (single-column working draft) + paper_overleaf_twocolumn.docx (two-column)
"""
import os
import re
import shutil
import subprocess
import sys

sys.path.insert(0, os.path.dirname(__file__))
from merge_tex_for_docx import expand

BASE = r"E:\论文\sci_redo\paper"
PY = r"E:\论文\期刊升级\.venv\Scripts\python.exe"
MERGED = os.path.join(BASE, "_merged_for_docx.tex")
DOCX = r"E:\论文\sci_redo\paper_overleaf.docx"
DOCX2 = r"E:\论文\sci_redo\paper_overleaf_twocolumn.docx"
BIB = os.path.join(BASE, "references.bib")
CSL = r"E:\论文\sci_redo\ieee.csl"


def preprocess(t):
    """Pandoc-friendly cleanup: table* -> table, remove resizebox/booktabs, keep captions."""
    t = t.replace(r"\begin{table*}", r"\begin{table}")
    t = t.replace(r"\end{table*}", r"\end{table}")
    t = re.sub(r"\\resizebox\{[^}]*\}\{[^}]*\}\{", "", t)
    t = re.sub(r"\}\s*\\end\{table\}", r"\\end{table}", t)
    t = t.replace("\\toprule", "").replace("\\midrule", "").replace("\\bottomrule", "")
    t = t.replace("\\centering", "")
    return t


def main():
    merged = expand(os.path.join(BASE, "main.tex"))
    merged = preprocess(merged)
    with open(MERGED, "w", encoding="utf-8") as f:
        f.write(merged)
    print("merged+preprocessed:", len(merged), "chars, tables:", merged.count("\\begin{table}"))

    import pypandoc
    pypandoc.convert_file(MERGED, "docx", format="latex", outputfile=DOCX,
        extra_args=["--citeproc", "--bibliography=" + BIB, "--csl=" + CSL,
                    "-M", "reference-section-title=References", "--wrap=none",
                    "--resource-path=" + BASE])
    print("docx OK")

    # ---- Typesetting polish: Times New Roman + remove blue heading color + two-column copy ----
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Pt

    for dst, twocol in [(DOCX, False), (DOCX2, True)]:
        if twocol:
            shutil.copy(DOCX, DOCX2)
            doc = Document(DOCX2)
        else:
            doc = Document(DOCX)
        # Body font
        for p in doc.paragraphs:
            for r in p.runs:
                r.font.name = "Times New Roman"
                if r.font.size is None:
                    r.font.size = Pt(10)
        # Table font
        for tb in doc.tables:
            for row in tb.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.name = "Times New Roman"
                            if r.font.size is None:
                                r.font.size = Pt(9)
        if twocol:
            for section in doc.sections:
                sectPr = section._sectPr
                cols = sectPr.find(qn("w:cols"))
                if cols is None:
                    cols = sectPr.makeelement(qn("w:cols"), {})
                    sectPr.append(cols)
                cols.set(qn("w:num"), "2")
                cols.set(qn("w:space"), "720")
                cols.set(qn("w:equalWidth"), "1")
        doc.save(dst)
    print("fonts + twocolumn copy done")


if __name__ == "__main__":
    main()
